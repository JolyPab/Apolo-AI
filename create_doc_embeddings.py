import os
import json
from dotenv import load_dotenv
from docx import Document
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from tqdm import tqdm
import base64

"""Скрипт создаёт FAISS-индекс из текста .docx-файла
Usage:
    python create_doc_embeddings.py <path_to_docx> [index_dir]
По умолчанию index_dir = "doc_faiss".
"""

import sys

# --- cli args ---------------------------------------------------------------
if len(sys.argv) < 2:
    print("Uso: python create_doc_embeddings.py <archivo.docx> [directorio_indice]")
    sys.exit(1)

docx_path = sys.argv[1]
index_dir = sys.argv[2] if len(sys.argv) > 2 else "doc_faiss"
images_dir = f"{index_dir}_images"

# --- env & embeddings -------------------------------------------------------
load_dotenv()
embeddings_model = AzureOpenAIEmbeddings(
    api_key=os.getenv("AZURE_EMBEDDINGS_API_KEY"),
    azure_endpoint=os.getenv("AZURE_EMBEDDINGS_ENDPOINT"),
    deployment="text-embedding-ada-002",
    api_version="2023-05-15",
)

# --- extract text -----------------------------------------------------------
print(f"📖 Extrayendo texto de {docx_path}…")
doc = Document(docx_path)
paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# --- extract text from tables ---------------------------------------------
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells if c.text.strip()]
        if cells:
            paragraphs.append(" | ".join(cells))

# --- extract images -------------------------------------------------------
print(f"🖼️ Extrayendo imágenes...")
os.makedirs(images_dir, exist_ok=True)
image_info = []

# Извлекаем изображения из relationships
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        try:
            image_data = rel.target_part.blob
            image_ext = rel.target_ref.split('.')[-1]
            image_filename = f"image_{len(image_info)+1}.{image_ext}"
            image_path = os.path.join(images_dir, image_filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            image_info.append({
                "filename": image_filename,
                "path": image_path,
                "size": len(image_data)
            })
            print(f"  ✅ {image_filename} ({len(image_data)} bytes)")
        except Exception as e:
            print(f"  ❌ Error extrayendo imagen: {e}")

# Сохраняем информацию об изображениях
with open(f"{images_dir}/images_info.json", "w", encoding="utf-8") as f:
    json.dump(image_info, f, ensure_ascii=False, indent=2)

print(f"💡 {len(paragraphs)} párrafos encontrados. Generando embeddings…")

# --- chunking ---------------------------------------------------------------
CHUNK_SIZE = 4  # абзацев в одном чанке
faiss_index = None
chunk = []

for para in tqdm(paragraphs, desc="Embeddings"):
    chunk.append(para)
    if len(chunk) >= CHUNK_SIZE:
        text = "\n".join(chunk)
        emb = FAISS.from_texts([text], embeddings_model)
        if faiss_index is None:
            faiss_index = emb
        else:
            faiss_index.merge_from(emb)
        chunk = []

# остаток
if chunk:
    text = "\n".join(chunk)
    emb = FAISS.from_texts([text], embeddings_model)
    if faiss_index is None:
        faiss_index = emb
    else:
        faiss_index.merge_from(emb)

# --- save -------------------------------------------------------------------
faiss_index.save_local(index_dir)
print(f"✅ Índice guardado en {index_dir}")
print(f"✅ {len(image_info)} imágenes guardadas en {images_dir}") 